import os
import re
import json
import logging
import datetime
import uuid
from pathlib import Path
import fitz  # PyMuPDF

from flask import Flask, render_template, request, flash, url_for, redirect, send_from_directory, session
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
import spacy
from docx import Document
from docx.shared import RGBColor
from langdetect import detect
from openai import OpenAI
from dotenv import load_dotenv
from flask import jsonify  # если ещё не импортировал
from time import time
from flask import request, jsonify, make_response

try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

# --- SETUP ---
load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
from flask import make_response, request
# --- FINAL WSGI-LEVEL CORS SANITIZER FOR /api/* ---
class CORSSanitizerMiddleware:
    def __init__(self, app):
        self.app = app

    def __call__(self, environ, start_response):
        path = (environ.get("PATH_INFO") or "")
        method = (environ.get("REQUEST_METHOD") or "GET").upper()

        # Short-circuit real preflights BEFORE Flask runs
        if path.startswith("/api/") and method == "OPTIONS":
            headers = [
                ("Access-Control-Allow-Origin", "*"),
                ("Access-Control-Allow-Methods", "GET, POST, OPTIONS"),
                ("Access-Control-Allow-Headers", "Content-Type"),
                ("Access-Control-Max-Age", "86400"),
                ("Vary", "Origin"),
                ("Cache-Control", "no-store"),
            ]
            start_response("204 No Content", headers)
            return [b""]

        # Otherwise, run the app and then sanitize headers
        def cors_start_response(status, headers, exc_info=None):
            if path.startswith("/api/"):
                # Remove any cookies/credentialed CORS that leaked in
                filtered = []
                for k, v in headers:
                    kl = k.lower()
                    if kl in ("set-cookie",
                              "access-control-allow-credentials",
                              "access-control-expose-headers"):
                        continue
                    filtered.append((k, v))
                headers = filtered

                # Enforce stateless CORS
                # (either set or overwrite existing)
                def put(name, val):
                    # overwrite if exists
                    for i, (k, _) in enumerate(headers):
                        if k.lower() == name.lower():
                            headers[i] = (name, val)
                            break
                    else:
                        headers.append((name, val))

                put("Access-Control-Allow-Origin", "*")
                put("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
                put("Access-Control-Allow-Headers", "Content-Type")
                put("Access-Control-Max-Age", "86400")
                put("Vary", "Origin")
                put("Cache-Control", "no-store")

                # If some auto handler set Allow on OPTIONS, we ignore it here.
                if method == "OPTIONS":
                    # Force a real preflight semantics if needed
                    status = "204 No Content"

            return start_response(status, headers, exc_info)

        return self.app(environ, cors_start_response)

# Wrap the Flask app
app.wsgi_app = CORSSanitizerMiddleware(app.wsgi_app)

@app.before_request
def short_circuit_api_options():
    # Return our own preflight before any other handler/middleware runs
    if request.method == "OPTIONS" and request.path.startswith("/api/"):
        resp = make_response("", 204)
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
        resp.headers["Access-Control-Max-Age"] = "86400"
        resp.headers["X-Parafix-Preflight"] = "1"
        return resp
from flask.sessions import SecureCookieSessionInterface

class ApiNoSessionInterface(SecureCookieSessionInterface):
    def save_session(self, app, session, response):
        # Never set a session cookie for API routes (OPTIONS/POST/GET)
        if request.path.startswith("/api/"):
            return
        return super().save_session(app, session, response)

app.session_interface = ApiNoSessionInterface()

# ---- Cookies & sessions (built-in Flask cookies, not Flask-Session) ----
app.config.update(
    SECRET_KEY=os.environ.get("SECRET_KEY", "change-me"),
    SESSION_COOKIE_SAMESITE="None",   # allow cross-site cookie (Lovable -> Render)
    SESSION_COOKIE_SECURE=True,       # required with SameSite=None (needs HTTPS)
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_NAME="parafix_sid",
    # REMOVE Flask-Session settings (we're not using that extension right now):
    # SESSION_TYPE="filesystem",
    # SESSION_FILE_DIR=os.path.join("/tmp", "flask_session"),
    # SESSION_PERMANENT=False,
    MAX_CONTENT_LENGTH=20 * 1024 * 1024,
    PREFERRED_URL_SCHEME="https",
)

app.config["SESSION_REFRESH_EACH_REQUEST"] = False

# ---- FINAL OVERRIDE: make all /api/* calls stateless + proper CORS ----
@app.after_request
def _final_api_cors_override(resp):
    try:
        p = request.path or ""
        if p.startswith("/api/"):
            # Always stateless CORS
            resp.headers["Access-Control-Allow-Origin"] = "*"
            resp.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
            resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
            resp.headers["Access-Control-Max-Age"] = "86400"
            # Strip any credentialed leftovers from other middlewares
            for h in ("Access-Control-Allow-Credentials", "Access-Control-Expose-Headers", "Set-Cookie"):
                resp.headers.pop(h, None)

            if request.method == "OPTIONS":
                # If Flask auto-OPTIONS already fired with 200, force it to be a real preflight
                resp.status_code = 204
                resp.set_data(b"")
                # Optional debug marker (helps confirm this ran)
                resp.headers["X-Parafix-Preflight"] = "1"
                # 'Allow' header is not needed for preflight
                resp.headers.pop("Allow", None)
    except Exception:
        pass
    return resp



@app.route("/api/analyze", methods=["OPTIONS"])
def analyze_preflight():
    resp = make_response("", 204)
    resp.headers["Access-Control-Allow-Origin"] = "*"
    resp.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
    resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
    resp.headers["Access-Control-Max-Age"] = "86400"
    resp.headers["X-Parafix-Preflight"] = "1"   # <— debug
    return resp

# --- File system paths ---
UPLOAD_FOLDER = "uploads"
STATIC_FOLDER = "static"
Path(UPLOAD_FOLDER).mkdir(parents=True, exist_ok=True)
Path(STATIC_FOLDER).mkdir(parents=True, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["STATIC_FOLDER"] = STATIC_FOLDER


client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# --- spaCy models (with graceful fallback) ---
def _load_spacy(model_name, lang_code):
    try:
        return spacy.load(model_name)
    except Exception as e:
        logger.warning(f"spaCy model '{model_name}' not found: {e}. Falling back to blank '{lang_code}' with sentencizer.")
        nlp = spacy.blank(lang_code)
        if "sentencizer" not in nlp.pipe_names:
            nlp.add_pipe("sentencizer")
        return nlp

nlp_en = _load_spacy("en_core_web_sm", "en")
nlp_ru = _load_spacy("ru_core_news_sm", "ru")

# --- LOGGING ---
def log_event(event, details=None):
    with open("analytics.log", "a") as f:
        f.write(json.dumps({
            "timestamp": str(datetime.datetime.now()),
            "event": event,
            "details": details or {}
        }) + "\n")

# --- CLEANUP ---
def cleanup_static_folder(days=1):
    now = datetime.datetime.now()
    for f in os.listdir(STATIC_FOLDER):
        path = os.path.join(STATIC_FOLDER, f)
        if os.path.isfile(path):
            mtime = datetime.datetime.fromtimestamp(os.path.getmtime(path))
            if (now - mtime).days > days:
                try:
                    os.remove(path)
                except Exception:
                    pass

# --- TEXT EXTRACTION ---
def extract_text_from_pdf(pdf_path):
    logger.info("Extracting text from PDF (PyMuPDF)")
    with fitz.open(pdf_path) as doc:
        return "\n".join(page.get_text("text") for page in doc).strip()
    

def extract_text_from_docx(docx_path):
    logger.info("Extracting text from DOCX")
    doc = Document(docx_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text]).strip()

# --- OPTIONAL: local NER (not used by default UI) ---
def extract_entities_spacy(text, lang='en'):
    nlp = nlp_en if lang == 'en' else nlp_ru
    doc = nlp(text)
    entities = []
    for ent in doc.ents:
        entity_type = None
        if ent.label_ in ["PERSON", "ORG"]:
            entity_type = "party"
        elif ent.label_ == "DATE":
            entity_type = "date"
        elif ent.label_ == "MONEY":
            entity_type = "monetary"
        elif ent.label_ in ["GPE", "LOC"]:
            entity_type = "place"
        if entity_type:
            clause = ent.sent.text if hasattr(ent, 'sent') else ent.text
            entities.append({"type": entity_type, "text": ent.text, "location": clause[:80]})
    return entities

# --- OpenAI helpers ---
def _extract_json(s: str):
    """Best-effort JSON extraction from model output."""
    s = s.strip()
    try:
        return json.loads(s)
    except Exception:
        pass
    # try object block
    start = s.find("{"); end = s.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(s[start:end+1])
        except Exception:
            pass
    # try array block
    start = s.find("["); end = s.rfind("]")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(s[start:end+1])
        except Exception:
            pass
    raise ValueError("Model did not return valid JSON")

def _chat_json(messages, model="gpt-4o"):
    resp = client.chat.completions.create(model=model, messages=messages)
    content = resp.choices[0].message.content or ""
    return _extract_json(content)

# --- Harvard partner analyzer + FAQ ---
def analyze_contract(text, language, is_premium=False):
    """
    Returns dict:
      {
        clauses: [{id, clause_type, risk_level, clause_quote, offset_start, offset_end,
                   why:[...], improved, flags:{unusual, one_sided, missing:[]}, evidence_quote}],
        alerts:  [{key, why, evidence_quote, suggested_wording}],
        faq:     [{question, answer, related_clause_type, risk_level, section_hint}],
        risk_count: int,
        risk_score: "low"|"medium"|"high"
      }
    """
    logger.info(f"Analyze contract (lang={language}, premium={is_premium})")

    # --- 1) Chunk politely using spaCy (consistent) ---
    nlp = nlp_en if language == 'en' else nlp_ru
    doc = nlp(text)
    sentences = [s.text for s in doc.sents if s.text.strip()]
    chunks, cur = [], ""
    for s in sentences:
        if len(cur) + len(s) <= 1400:
            cur += (s + " ")
        else:
            chunks.append(cur.strip())
            cur = s + " "
    if cur:
        chunks.append(cur.strip())

    system_msg = {
        "role": "system",
        "content": (
            "You are a Harvard Law–trained senior partner and former prosecutor, expert in commercial contracts. "
            "Be precise, practical, and concise. "
            f"Detected language: {language}. Respond ONLY in this language. "
            "Return ONLY valid JSON that matches the schema—no extra text."
        )
    }

    schema_instr = (
        "Output a JSON object with keys:\n"
        "  clauses: array of objects with keys:\n"
        "    - id: string\n"
        "    - clause_type: string (e.g., indemnity, liability, termination, payment, ip, confidentiality, governing_law, dispute, data_protection, reps, warranties)\n"
        "    - risk_level: one of low|medium|high\n"
        "    - clause_quote: string (EXACT substring from Context)\n"
        "    - offset_start: integer (0-based index within Context)\n"
        "    - offset_end: integer\n"
        "    - why: array of 3–6 short bullets in plain language\n"
        "    - improved: string (professionally improved wording, balanced by default)\n"
        "    - flags: object { unusual: boolean, one_sided: boolean, missing: array of strings }\n"
        "    - evidence_quote: short quote (<=30 words) from Context supporting the risk\n"
        "  alerts: array of objects with keys:\n"
        "    - key: machine key (e.g., termination_missing, unlimited_liability, unilateral_indemnity)\n"
        "    - why: short explanation\n"
        "    - evidence_quote: short quote\n"
        "    - suggested_wording: concise standard wording to resolve it\n"
        "Do NOT include FAQ here."
    )

    all_clauses, all_alerts = [], []

    for idx, chunk in enumerate(chunks):
        user_msg = {
            "role": "user",
            "content": (
                "Analyze the following contract Context and return the JSON per schema.\n\n"
                f"{schema_instr}\n\n"
                "Context:\n<<<" + chunk + ">>>"
            )
        }
        try:
            obj = _chat_json([system_msg, user_msg])
        except Exception as e:
            logger.warning(f"Chunk {idx} JSON parse failed: {e}")
            continue

        clauses = obj.get("clauses", []) if isinstance(obj, dict) else []
        alerts  = obj.get("alerts", [])  if isinstance(obj, dict) else []
        all_clauses.extend(clauses)
        all_alerts.extend(alerts)

    # De-dup clauses by (quote, type)
    seen = set()
    deduped = []
    for c in all_clauses:
        key = (c.get("clause_quote","").strip(), c.get("clause_type",""))
        if key in seen:
            continue
        seen.add(key)
        deduped.append(c)

    # Risk score (max severity present)
    levels = [c.get("risk_level","low") for c in deduped]
    if "high" in levels:
        risk_score = "high"
    elif "medium" in levels:
        risk_score = "medium"
    else:
        risk_score = "low"

    # FAQ for whole doc
    try:
        faq = generate_faq(full_text=text, language=language, num_q=10)
    except Exception as e:
        logger.warning(f"FAQ generation failed: {e}")
        faq = []

    return {
        "clauses": deduped,
        "alerts": all_alerts,
        "faq": faq,
        "risk_count": len(deduped),
        "risk_score": risk_score
    }

def generate_faq(full_text: str, language: str, num_q=10):
    """
    Harvard-level Q&A for the whole contract.
    Returns list[{question, answer, related_clause_type, risk_level, section_hint}]
    """
    system_msg = {
        "role": "system",
        "content": (
            "You are a Harvard Law–trained senior partner and former prosecutor. "
            "Craft incisive, practical Q&A a client can act on immediately. "
            f"Detected language: {language}. Respond ONLY in this language. "
            "Return ONLY valid JSON per schema."
        )
    }
    user_msg = {
        "role": "user",
        "content": (
            "Read the full contract text and create high-value FAQs.\n"
            "Output JSON with key 'faq' = array of objects with:\n"
            "  - question: crisp, client-facing question (no legalese)\n"
            "  - answer: concise, practical answer (2–5 sentences), include conditions/exceptions\n"
            "  - related_clause_type: e.g., termination, liability, indemnity, payment, confidentiality, ip, data_protection, governing_law, dispute\n"
            "  - risk_level: low|medium|high (exposure if misunderstood)\n"
            "  - section_hint: best-effort section/heading; null if unknown\n"
            f"Produce {num_q} Q&A pairs.\n\n"
            "Full text:\n<<<" + full_text[:120000] + ">>>"
        )
    }
    obj = _chat_json([system_msg, user_msg])
    return obj.get("faq", []) if isinstance(obj, dict) else []

# --- DOCX REPORT (clean, non-redline) ---
def generate_report_docx(analysis, language: str):
    """
    Build a readable report:
      Summary -> Alerts -> Risky Clauses -> FAQ
    """
    doc = Document()
    title = "Parafix Contract Review Report" if language == "en" else "Отчет Parafix по проверке договора"
    doc.add_heading(title, level=1)

    # Summary
    doc.add_heading("Summary" if language == "en" else "Итоги", level=2)
    p = doc.add_paragraph()
    p.add_run(("Risk score: " if language == "en" else "Уровень риска: ")).bold = True
    p.add_run(analysis.get("risk_score", "low").upper())

    p2 = doc.add_paragraph()
    p2.add_run(("Risky clauses found: " if language == "en" else "Найдено рискованных пунктов: ")).bold = True
    p2.add_run(str(analysis.get("risk_count", 0)))

    # Alerts
    alerts = analysis.get("alerts", [])
    if alerts:
        doc.add_heading("Alerts" if language == "en" else "Предупреждения", level=2)
        for a in alerts:
            doc.add_heading(a.get("key","alert"), level=3)
            doc.add_paragraph(a.get("why",""))
            ev = a.get("evidence_quote")
            if ev:
                doc.add_paragraph(("Evidence: " if language == "en" else "Доказательство: ") + ev)
            sw = a.get("suggested_wording")
            if sw:
                run = doc.add_paragraph().add_run(("Suggested wording: " if language == "en" else "Рекомендуемая формулировка: "))
                run.bold = True
                doc.add_paragraph(sw)

    # Risky Clauses
    clauses = analysis.get("clauses", [])
    if clauses:
        doc.add_heading("Risky Clauses" if language == "en" else "Рискованные пункты", level=2)
        for c in clauses:
            hdr = f"[{c.get('risk_level','low').upper()}] {c.get('clause_type','clause')}"
            doc.add_heading(hdr, level=3)
            doc.add_paragraph(("Original:" if language == "en" else "Оригинал:"))
            q = doc.add_paragraph(c.get("clause_quote",""))
            q.runs[0].font.color.rgb = RGBColor(128, 0, 0)

            doc.add_paragraph(("Why risky:" if language == "en" else "Почему рискованно:"))
            for b in c.get("why", []):
                doc.add_paragraph(b, style='List Bullet')

            doc.add_paragraph(("Improved clause:" if language == "en" else "Улучшенная редакция:"))
            imp = doc.add_paragraph(c.get("improved",""))
            imp.runs[0].font.color.rgb = RGBColor(0, 100, 0)

    # FAQ
    faq = analysis.get("faq", [])
    if faq:
        doc.add_heading("FAQ" if language == "en" else "Вопросы и ответы", level=2)
        for qa in faq:
            q = qa.get("question","Q")
            a = qa.get("answer","A")
            doc.add_paragraph(("Q: " if language == "en" else "В: ") + q)
            doc.add_paragraph(("A: " if language == "en" else "О: ") + a)

    footer = doc.add_paragraph()
    footer.add_run(
        f"Generated by Parafix on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        if language == "en"
        else f"Сформировано Parafix {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ).italic = True
    return doc

def convert_docx_to_pdf(docx_path, pdf_path):
    if DOCX2PDF_AVAILABLE:
        try:
            convert(docx_path, pdf_path)
            return True
        except Exception as e:
            logger.warning(f"docx2pdf failed: {e}")
            return False
    return False

# --- ROUTES ---
@app.route('/login', methods=['GET', 'POST'])
def login():
    # Lovable JSON login (email + optional license_key)
    if request.method == 'POST' and request.form.get('email'):
        email = request.form.get('email', '').strip()
        license_key = request.form.get('license_key', '').strip()
        if license_key == "premium":
            session['is_premium'] = True
        session['user'] = email
        return jsonify({
            "ok": True,
            "email": email,
            "is_premium": bool(session.get('is_premium'))
        }), 200

    # Your existing HTML login (username/password) stays intact
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        if username == "premium" and password == "premium":
            session['is_premium'] = True
            flash("Premium login successful.", "success")
            return redirect(url_for('index'))
        flash("Invalid credentials.", "danger")
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('is_premium', None)
    flash("Logged out.", "info")
    return redirect(url_for('index'))

@app.route('/static/<path:filename>')
def static_files(filename):
    return send_from_directory(app.config['STATIC_FOLDER'], filename, as_attachment=True)

@app.route('/', methods=['GET', 'POST'])
def index():
    cleanup_static_folder(days=1)
    if request.method == 'POST':
        file = request.files.get('file')
        if not file or file.filename == '':
            flash("No file uploaded", "danger")
            return redirect(url_for('index'))

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            # Detect file type & extract
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(filepath)
            elif filename.lower().endswith('.docx'):
                text = extract_text_from_docx(filepath)
            else:
                flash('Please upload a PDF or DOCX file.', 'danger')
                os.remove(filepath)
                return redirect(url_for('index'))

            if not text.strip():
                flash("No text extracted from file.", "danger")
                os.remove(filepath)
                return redirect(url_for('index'))

            # Prevent re-uploads of generated report
            if "Generated by Parafix" in text or "Improved by ContractAI Pro" in text:
                flash("This document appears to be a previously generated report. Please upload an original contract.", "warning")
                os.remove(filepath)
                return redirect(url_for('index'))

            # Language detect -> RU/EN default
            language = detect(text) if text.strip() else 'en'
            if language not in ['en', 'ru']:
                language = 'en'

            is_premium = session.get('is_premium', False)

            # --- Analyze (Harvard-level + FAQ) ---
            analysis = analyze_contract(text, language, is_premium=is_premium)

            # Entities: only for premium (as before)
            entities = None

            # --- Build .docx Report ---
            report_doc = generate_report_docx(analysis, language)
            base_name = os.path.splitext(filename)[0]
            timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
            unique_id = str(uuid.uuid4())[:8]
            report_name = f"parafix_report_{base_name}_{timestamp}_{unique_id}.docx"
            report_path = os.path.join(STATIC_FOLDER, report_name)
            report_doc.save(report_path)

            # Analytics
            log_event("contract_upload", {
                "filename": filename,
                "language": language,
                "risk_count": analysis.get("risk_count", 0),
                "risk_score": analysis.get("risk_score", "low"),
                "is_premium": is_premium
            })

            # Clean uploaded file
            if os.path.exists(filepath):
                os.remove(filepath)

            return render_template(
                'results.html',
                analysis=analysis,           # now includes clauses, alerts, faq, risk_score
                docx_path=report_name,       # small download button can use this
                language=language,
                entities=entities,
                is_premium=is_premium
            )

        except Exception as e:
            logger.error(f"Processing error: {e}")
            if os.path.exists(filepath):
                os.remove(filepath)
            flash(f"Error processing contract: {str(e)}", 'danger')
            return redirect(url_for('index'))

    return render_template('index.html', is_premium=session.get('is_premium', False))
# --- JSON API for Lovable (no templates, just JSON) ---
@app.route("/api/analyze", methods=["POST"], provide_automatic_options=False)
def api_analyze():
    cleanup_static_folder(days=1)

    # Stateless MVP: no session/auth here; wire JWT or header later if needed
    is_premium = False

    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({"error": "No file uploaded"}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    try:
        # Extract text (PDF / DOCX)
        if filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(filepath)
        elif filename.lower().endswith('.docx'):
            text = extract_text_from_docx(filepath)
        else:
            os.remove(filepath)
            return jsonify({"error": "Please upload a PDF or DOCX file."}), 400

        if not text.strip():
            os.remove(filepath)
            return jsonify({"error": "No text extracted from file."}), 400

        # Block re-uploads of generated reports
        if "Generated by Parafix" in text or "Improved by ContractAI Pro" in text:
            os.remove(filepath)
            return jsonify({"error": "This looks like a generated report. Please upload the original contract."}), 400

        # Language detect
        language = detect(text) if text.strip() else 'en'
        if language not in ['en', 'ru']:
            language = 'en'

        # Analyze + Report
        analysis = analyze_contract(text, language, is_premium=is_premium)
        report_doc = generate_report_docx(analysis, language)
        base_name = os.path.splitext(filename)[0]
        timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
        unique_id = str(uuid.uuid4())[:8]
        report_name = f"parafix_report_{base_name}_{timestamp}_{unique_id}.docx"
        report_path = os.path.join(STATIC_FOLDER, report_name)
        report_doc.save(report_path)

        # Analytics
        log_event("contract_upload", {
            "filename": filename,
            "language": language,
            "risk_count": analysis.get("risk_count", 0),
            "risk_score": analysis.get("risk_score", "low"),
            "is_premium": is_premium
        })

        # Clean uploaded file
        if os.path.exists(filepath):
            os.remove(filepath)

        # Save analysis JSON to static (client can fetch/download)
        analysis_id = str(uuid.uuid4())[:8]
        json_name = f"analysis_{analysis_id}.json"
        json_path = os.path.join(STATIC_FOLDER, json_name)
        with open(json_path, "w", encoding="utf-8") as jf:
            json.dump(analysis, jf, ensure_ascii=False)

        # Return JSON — CORS headers will be added by @after_request
        return jsonify({
            "analysis": analysis,
            "docx_path": report_name,
            "language": language,
            "analysis_id": analysis_id
        }), 200

    except Exception as e:
        logger.error(f"Processing error: {e}")
        if os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({"error": f"Error processing contract: {str(e)}"}), 500



@app.get("/api/analysis/latest")
def api_analysis_latest():
    analysis = {}
    analysis_id = session.get("analysis_id")
    if analysis_id:
        path = os.path.join(STATIC_FOLDER, f"analysis_{analysis_id}.json")
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as jf:
                analysis = json.load(jf)

    return jsonify({
        "analysis": analysis,
        "docx_path": session.get("docx_path", ""),
        "language": session.get("language", "en"),
        "analysis_id": analysis_id
    }), 200


@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(error):
    flash("File is too large! Maximum upload size is 20 MB.", "danger")
    return redirect(url_for('index'))

@app.route("/health", methods=["GET"])
def health_root():
    return jsonify(ok=True, ts=int(time()))

@app.route("/api/health", methods=["GET"])
def health_api():
    return jsonify(ok=True, ts=int(time()))

# --- Debug: list all routes live ---
@app.get("/__routes")
def list_routes():
    return jsonify(sorted(str(r) for r in app.url_map.iter_rules()))

# --- Debug: show running commit and routes ---
BUILD = os.getenv("RENDER_GIT_COMMIT", "dev")[:7]

@app.get("/version")
def version():
    return jsonify(
        commit=BUILD,
        routes=sorted(r.rule for r in app.url_map.iter_rules())
    )
       


if __name__ == '__main__':
    app.run(debug=True)
